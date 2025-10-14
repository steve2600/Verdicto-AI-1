"""
Enhanced RAG utils for insurance documents - Semantic Similarity approach
Supports PDF, DOCX, and Email processing with semantic similarity-based chunking
"""

import os
import re
import json
import tempfile
import asyncio
import traceback
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from collections import defaultdict
import httpx
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders.word_document import Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from starlette.concurrency import run_in_threadpool

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx2txt
except ImportError:
    docx2txt = None

try:
    from python_docx import Document as DocxDocument
except ImportError:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        DocxDocument = None

# Environment setup
VOYAGE_API_KEY = os.getenv("VOYAGE_API_KEY")
if not VOYAGE_API_KEY:
    raise ValueError("VOYAGE_API_KEY not found in environment")


def extract_pdf_preview(file_path: str) -> str:
    """Extract first 50 words from PDF for preview"""
    return "PDF preview disabled"


def extract_docx_preview(file_path: str) -> str:
    """Extract first 50 words from DOCX for preview"""
    try:
        if docx2txt:
            # Try with docx2txt first (simpler)
            text = docx2txt.process(file_path)
            if text:
                words = text.strip().split()
                preview_words = words[:50] if len(words) >= 50 else words
                return ' '.join(preview_words) + ('...' if len(words) > 50 else '')

        if DocxDocument:
            # Fallback to python-docx
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs[:5]:  # First 5 paragraphs
                text += paragraph.text + " "

            words = text.strip().split()
            preview_words = words[:50] if len(words) >= 50 else words
            return ' '.join(preview_words) + ('...' if len(words) > 50 else '')

        return "DOCX processing libraries not available"

    except Exception as e:
        return f"Could not extract preview: {str(e)}"


def extract_email_preview(email_content: str) -> str:
    """Extract preview from email content"""
    try:
        # Parse email
        msg = email.message_from_string(email_content)

        # Get subject and sender info
        subject = msg.get('subject', 'No Subject')
        sender = msg.get('from', 'Unknown Sender')

        # Get body content
        body = ""
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    body = part.get_payload(decode=True)
                    if isinstance(body, bytes):
                        body = body.decode('utf-8', errors='ignore')
                    break
        else:
            body = msg.get_payload(decode=True)
            if isinstance(body, bytes):
                body = body.decode('utf-8', errors='ignore')

        # Create preview
        preview_text = f"From: {sender} | Subject: {subject}"
        if body:
            body_words = body.strip().split()[:30]  # First 30 words of body
            if body_words:
                preview_text += f" | Content: {' '.join(body_words)}"
                if len(body.split()) > 30:
                    preview_text += "..."

        return preview_text

    except Exception as e:
        return f"Could not extract email preview: {str(e)}"


def detect_file_type(file_path: str) -> str:
    """Detect file type based on extension and content"""
    file_path = Path(file_path)
    extension = file_path.suffix.lower()

    if extension == '.pdf':
        return 'pdf'
    elif extension in ['.docx', '.doc']:
        return 'docx'
    elif extension in ['.eml', '.msg']:
        return 'email'
    elif extension == '.txt':
        # Check if it's an email file
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read(1000)  # Read first 1000 chars
                if any(header in content.lower() for header in ['from:', 'to:', 'subject:', 'date:']):
                    return 'email'
        except:
            pass
        return 'text'
    else:
        # Try to detect based on content
        mime_type, _ = mimetypes.guess_type(str(file_path))
        if mime_type:
            if 'pdf' in mime_type:
                return 'pdf'
            elif 'word' in mime_type or 'officedocument' in mime_type:
                return 'docx'
            elif 'text' in mime_type:
                return 'text'

        return 'unknown'


async def get_embeddings(texts: List[str], model: str = "voyage-3-large") -> List[List[float]]:
    """Get embeddings from Voyage AI API"""
    if not VOYAGE_API_KEY:
        raise ValueError("VOYAGE_API_KEY not found in environment")

    async with httpx.AsyncClient(timeout=30.0) as client:
        response = await client.post(
            "https://api.voyageai.com/v1/embeddings",
            headers={
                "Authorization": f"Bearer {VOYAGE_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "input": texts,
                "model": model
            }
        )
        response.raise_for_status()
        data = response.json()
        return [item["embedding"] for item in data]


def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
    """Calculate cosine similarity between two vectors"""
    vec1 = np.array(vec1)
    vec2 = np.array(vec2)

    dot_product = np.dot(vec1, vec2)
    norm1 = np.linalg.norm(vec1)
    norm2 = np.linalg.norm(vec2)

    if norm1 == 0 or norm2 == 0:
        return 0.0

    return dot_product / (norm1 * norm2)


async def calculate_semantic_scores(chunks: List[Document], query: str = None) -> List[float]:
    """Calculate semantic similarity scores for chunks"""

    # If no query provided, use a generic business/document query
    if not query:
        query = "important terms conditions coverage benefits requirements procedures definitions"

    # Prepare texts for embedding
    chunk_texts = []
    for chunk in chunks:
        # Use first 500 chars to avoid token limits while preserving context
        text = chunk.page_content[:500].strip()
        if len(text) < 50:  # If too short, use full text
            text = chunk.page_content.strip()
        chunk_texts.append(text)

    # Get embeddings for all chunks and query
    all_texts = [query] + chunk_texts

    try:
        embeddings = await get_embeddings(all_texts)
        query_embedding = embeddings[0]
        chunk_embeddings = embeddings[1:]

        # Calculate similarity scores
        scores = []
        for chunk_embedding in chunk_embeddings:
            similarity = cosine_similarity(query_embedding, chunk_embedding)
            scores.append(max(similarity, 0.0))  # Ensure non-negative

        return scores

    except Exception as e:
        # Fallback to basic text length scoring
        return [min(len(text.split()) / 100, 1.0) for text in chunk_texts]


async def semantic_chunk_sampling(chunks: List[Document], max_chunks: int = 100, query: str = None) -> List[Document]:
    """Smart sampling based on semantic similarity"""

    if len(chunks) <= max_chunks:
        return chunks

    # Get semantic similarity scores
    semantic_scores = await calculate_semantic_scores(chunks, query)

    # Combine chunks with their scores
    scored_chunks = list(zip(chunks, semantic_scores))

    # Group by section if available
    section_groups = defaultdict(list)
    standalone_chunks = []

    for chunk, score in scored_chunks:
        section_title = chunk.metadata.get('section_title', '')
        if section_title:
            section_groups[section_title].append((chunk, score))
        else:
            standalone_chunks.append((chunk, score))

    selected_chunks = []

    # Process sections - take top chunks from each section based on semantic score
    if section_groups:
        # Calculate average semantic score per section
        section_avg_scores = []
        for section_title, section_chunk_scores in section_groups.items():
            avg_score = sum(score for _, score in section_chunk_scores) / len(section_chunk_scores)
            section_avg_scores.append((section_title, avg_score, section_chunk_scores))

        # Sort sections by average semantic relevance
        section_avg_scores.sort(key=lambda x: x[1], reverse=True)

        # Distribute chunks across semantically relevant sections
        chunks_per_section = max(2, max_chunks // max(len(section_groups), 1))

        for section_title, avg_score, section_chunk_scores in section_avg_scores:
            # Sort chunks within section by semantic score
            section_chunk_scores.sort(key=lambda x: x[1], reverse=True)

            # Take top semantically similar chunks from this section
            section_limit = min(len(section_chunk_scores), chunks_per_section)
            selected_chunks.extend([chunk for chunk, score in section_chunk_scores[:section_limit]])

            if len(selected_chunks) >= max_chunks:
                break

    # Add standalone chunks based on semantic relevance
    remaining_slots = max_chunks - len(selected_chunks)
    if remaining_slots > 0 and standalone_chunks:
        standalone_chunks.sort(key=lambda x: x[1], reverse=True)
        selected_chunks.extend([chunk for chunk, score in standalone_chunks[:remaining_slots]])

    # Final selection if still over limit - keep highest semantic scores
    if len(selected_chunks) > max_chunks:
        # Re-score if needed
        final_scores = []
        for chunk in selected_chunks:
            # Find the chunk's semantic score from our original scoring
            for orig_chunk, score in scored_chunks:
                if orig_chunk == chunk:
                    final_scores.append((chunk, score))
                    break

        final_scores.sort(key=lambda x: x[1], reverse=True)
        selected_chunks = [chunk for chunk, score in final_scores[:max_chunks]]

    return selected_chunks


class UniversalDocumentSplitter:
    """Universal splitter that handles PDF, DOCX, and Email documents"""

    # Common section patterns for different document types
    SECTION_PATTERNS = {
        'insurance': [
            r'^\s*(?:SECTION\s+)?([IVXLCDM]+|\d+)\.\s*([A-Z][A-Z\s&-]+)(?:\s*[-–—]\s*)?',
            r'^\s*([A-Z][A-Z\s&-]{10,})\s*$',
            r'^\s*Part\s+([IVXLCDM]+|\d+)[\s:]?\s*([A-Z][A-Za-z\s&-]+)',
            r'^\s*Article\s+(\d+)[\s:]?\s*([A-Z][A-Za-z\s&-]+)',
            r'^\s*Chapter\s+(\d+)[\s:]?\s*([A-Z][A-Za-z\s&-]+)',
        ],
        'email': [
            r'^(From|To|Cc|Bcc|Subject|Date):\s*(.+)',
            r'^[-=_]{3,}\s*$',  # Email separators
            r'^>\s*(.+)',  # Quoted text
        ],
        'general': [
            r'^\s*([A-Z][A-Za-z\s&-]{5,})\s*$',  # General headers
            r'^\s*\d+\.\s*([A-Z][A-Za-z\s&-]+)',  # Numbered sections
        ]
    }

    def __init__(self, chunk_size: int = 2500, chunk_overlap: int = 400, doc_type: str = 'general'):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.doc_type = doc_type

    def split_documents(self, documents: List[Document]) -> List[Document]:
        """Split documents while preserving section boundaries"""
        all_chunks = []

        for doc in documents:
            chunks = self._split_single_document(doc)
            all_chunks.extend(chunks)

        return all_chunks

    def _split_single_document(self, document: Document) -> List[Document]:
        """Split a single document preserving sections based on document type"""
        text = document.page_content
        doc_type = document.metadata.get('doc_type', self.doc_type)

        if doc_type == 'email':
            return self._split_email_document(document)
        else:
            sections = self._identify_sections(text, doc_type)

            if not sections:
                return self._fallback_split(document)

            chunks = []
            for section_title, section_text, start_pos in sections:
                section_chunks = self._split_section(
                    section_text,
                    section_title,
                    document.metadata,
                    start_pos
                )
                chunks.extend(section_chunks)

            return chunks

    def _split_email_document(self, document: Document) -> List[Document]:
        """Special handling for email documents"""
        text = document.page_content

        # Parse email structure
        lines = text.split('\n')
        headers = {}
        body_start = 0

        # Extract headers
        for i, line in enumerate(lines):
            if ':' in line and i < 20:  # Check first 20 lines for headers
                key, value = line.split(':', 1)
                if key.strip().lower() in ['from', 'to', 'subject', 'date', 'cc', 'bcc']:
                    headers[key.strip().lower()] = value.strip()
                    body_start = i + 1

        # Get body
        body_lines = lines[body_start:]
        body = '\n'.join(body_lines)

        # Create chunks
        chunks = []

        # Header chunk
        if headers:
            header_text = "Email Headers:\n"
            for key, value in headers.items():
                header_text += f"{key.title()}: {value}\n"

            metadata = document.metadata.copy()
            metadata.update({
                'section_title': 'Email Headers',
                'chunk_type': 'email_headers',
                **headers
            })
            chunks.append(Document(page_content=header_text, metadata=metadata))

        # Body chunks
        if body.strip():
            if len(body) <= self.chunk_size:
                # Small body - keep as one chunk
                metadata = document.metadata.copy()
                metadata.update({
                    'section_title': 'Email Body',
                    'chunk_type': 'email_body',
                    **headers
                })
                chunks.append(Document(page_content=body, metadata=metadata))
            else:
                # Split large body
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=self.chunk_size,
                    chunk_overlap=self.chunk_overlap
                )
                body_chunks = splitter.split_text(body)

                for i, chunk_text in enumerate(body_chunks):
                    metadata = document.metadata.copy()
                    metadata.update({
                        'section_title': f'Email Body Part {i+1}',
                        'chunk_type': 'email_body',
                        'chunk_index': i,
                        **headers
                    })
                    chunks.append(Document(page_content=chunk_text, metadata=metadata))

        return chunks

    def _identify_sections(self, text: str, doc_type: str) -> List[tuple]:
        """Identify logical sections based on document type"""
        patterns = self.SECTION_PATTERNS.get(doc_type, self.SECTION_PATTERNS['general'])

        sections = []
        lines = text.split('\n')
        current_section = []
        current_title = ""
        section_start = 0

        for i, line in enumerate(lines):
            is_header = False
            title = ""

            for pattern in patterns:
                match = re.match(pattern, line.strip())
                if match:
                    is_header = True
                    title = line.strip()
                    break

            if is_header and current_section:
                section_text = '\n'.join(current_section)
                if len(section_text.strip()) > 50:
                    sections.append((current_title, section_text, section_start))

                current_section = [line]
                current_title = title
                section_start = i
            else:
                current_section.append(line)

        # Add final section
        if current_section:
            section_text = '\n'.join(current_section)
            if len(section_text.strip()) > 50:
                sections.append((current_title, section_text, section_start))

        return sections

    def _split_section(self, section_text: str, section_title: str,
                      base_metadata: Dict, start_pos: int) -> List[Document]:
        """Split a section while keeping related content together"""

        # For small sections, keep them whole
        if len(section_text) <= self.chunk_size * 1.5:
            metadata = base_metadata.copy()
            metadata.update({
                'section_title': section_title,
                'section_start': start_pos,
                'chunk_type': 'complete_section'
            })
            enhanced_content = f"{section_title}\n\n{section_text}" if section_title else section_text
            return [Document(page_content=enhanced_content, metadata=metadata)]

        # Split larger sections
        separators = [
            r'\n\s*(?:\([a-z]\)|\([0-9]+\)|\([IVXLCDM]+\))',
            r'\n\s*[a-z]\.\s+',
            r'\n\s*\d+\.\s+',
            r'(?<=[.!?])\s+(?=[A-Z])',
            r'\n\n+',
            r'\n',
            r'\.\s+',
            r'\s+'
        ]

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=separators,
            length_function=len,
            is_separator_regex=True,
        )

        section_chunks = splitter.split_text(section_text)
        chunks = []

        for i, chunk_text in enumerate(section_chunks):
            if self._has_sufficient_context(chunk_text):
                metadata = base_metadata.copy()
                metadata.update({
                    'section_title': section_title,
                    'section_start': start_pos,
                    'chunk_index': i,
                    'chunk_type': 'section_part'
                })
                enhanced_content = f"{section_title}\n\n{chunk_text}" if section_title else chunk_text
                chunks.append(Document(page_content=enhanced_content, metadata=metadata))

        return chunks

    def _has_sufficient_context(self, text: str) -> bool:
        """Check if chunk has sufficient context"""
        return len(text.strip()) > 10 and len(text.split()) > 2

    def _fallback_split(self, document: Document) -> List[Document]:
        """Fallback splitting when no sections are detected"""
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap
        )

        chunks = splitter.split_documents([document])

        filtered_chunks = []
        for i, chunk in enumerate(chunks):
            if self._has_sufficient_context(chunk.page_content):
                chunk.metadata.update({
                    'chunk_index': i,
                    'chunk_type': 'fallback_split'
                })
                filtered_chunks.append(chunk)

        return filtered_chunks


def load_docx_document(file_path: str) -> List[Document]:
    """Load DOCX document using available libraries"""
    try:
        if docx2txt:
            # Try docx2txt first (simpler and more reliable)
            text = docx2txt.process(file_path)
            if text and text.strip():
                metadata = {
                    'source': file_path,
                    'doc_type': 'docx',
                    'loader': 'docx2txt'
                }
                return [Document(page_content=text, metadata=metadata)]

        # Try Langchain's Docx2txtLoader
        try:
            loader = Docx2txtLoader(file_path)
            docs = loader.load()
            if docs:
                for doc in docs:
                    doc.metadata.update({'doc_type': 'docx', 'loader': 'langchain_docx2txt'})
                return docs
        except Exception:
            pass

        # Fallback to python-docx if available
        if DocxDocument:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            if text.strip():
                metadata = {
                    'source': file_path,
                    'doc_type': 'docx',
                    'loader': 'python_docx'
                }
                return [Document(page_content=text, metadata=metadata)]

        raise Exception("No DOCX processing library available")

    except Exception as e:
        raise Exception(f"Failed to load DOCX: {str(e)}")


def load_email_document(file_path: str) -> List[Document]:
    """Load email document (.eml, .msg, or text file with email content)"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        # Try to parse as email
        try:
            msg = email.message_from_string(content)

            # Extract metadata
            metadata = {
                'source': file_path,
                'doc_type': 'email',
                'subject': msg.get('subject', ''),
                'from': msg.get('from', ''),
                'to': msg.get('to', ''),
                'date': msg.get('date', ''),
                'loader': 'email_parser'
            }

            # Get full email content (headers + body)
            full_content = content

            return [Document(page_content=full_content, metadata=metadata)]

        except Exception:
            # If email parsing fails, treat as plain text with email markers
            if any(marker in content.lower() for marker in ['from:', 'to:', 'subject:']):
                metadata = {
                    'source': file_path,
                    'doc_type': 'email',
                    'loader': 'text_email'
                }
                return [Document(page_content=content, metadata=metadata)]
            else:
                raise Exception("File doesn't appear to contain email content")

    except Exception as e:
        raise Exception(f"Failed to load email: {str(e)}")


def get_smart_splitter(content_length: int, doc_type: str = 'general') -> UniversalDocumentSplitter:
    """Get document splitter with size-based parameters"""
    if content_length > 100000:  # Very large documents
        chunk_size, chunk_overlap = 3500, 500
    elif content_length > 50000:
        chunk_size, chunk_overlap = 3000, 400
    elif content_length > 25000:
        chunk_size, chunk_overlap = 2500, 350
    elif content_length > 10000:
        chunk_size, chunk_overlap = 2000, 300
    elif content_length > 5000:
        chunk_size, chunk_overlap = 1800, 250
    else:
        chunk_size, chunk_overlap = 1500, 200

    return UniversalDocumentSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        doc_type=doc_type
    )


async def load_document_semantic(path_or_url: str, max_chunks: int = 100, query: str = None) -> List[Document]:
    """Universal document loading with semantic similarity-based chunking"""
    temp_download = False
    file_path = None

    try:
        # Handle URL downloads
        if path_or_url.startswith(("http://", "https://")):
            print(f"Processing document from URL: {path_or_url}")

            async with httpx.AsyncClient(
                timeout=httpx.Timeout(60.0),
                limits=httpx.Limits(max_connections=1, max_keepalive_connections=0)
            ) as client:
                response = await client.get(path_or_url, follow_redirects=True)
                response.raise_for_status()

                # Determine file extension from URL or Content-Type
                content_type = response.headers.get('content-type', '').lower()
                if 'pdf' in content_type:
                    extension = '.pdf'
                elif 'word' in content_type or 'officedocument' in content_type:
                    extension = '.docx'
                else:
                    # Try to get from URL
                    url_path = Path(path_or_url)
                    extension = url_path.suffix or '.txt'

                file_path = Path(tempfile.mkstemp(suffix=extension)[1])
                with open(file_path, "wb") as f:
                    f.write(response.content)
                temp_download = True
        else:
            file_path = Path(path_or_url)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

        # Detect file type
        file_type = detect_file_type(str(file_path))
        print(f"Detected file type: {file_type}")

        def load_and_split():
            try:
                # Load document based on type
                if file_type == 'pdf':
                    loader = PyPDFLoader(str(file_path))
                    raw_docs = loader.load()
                    doc_type = 'insurance'  # Assume insurance for PDFs

                    # Extract and print preview
                    preview = extract_pdf_preview(str(file_path))
                    print(f"Document preview: {preview}")

                elif file_type == 'docx':
                    raw_docs = load_docx_document(str(file_path))
                    doc_type = 'general'

                    # Extract and print preview
                    preview = extract_docx_preview(str(file_path))
                    print(f"Document preview: {preview}")

                elif file_type == 'email':
                    raw_docs = load_email_document(str(file_path))
                    doc_type = 'email'

                    # Extract and print preview
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        email_content = f.read()
                    preview = extract_email_preview(email_content)
                    print(f"Email preview: {preview}")

                else:
                    # Fallback to text loading
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()

                    metadata = {'source': str(file_path), 'doc_type': 'text'}
                    raw_docs = [Document(page_content=content, metadata=metadata)]
                    doc_type = 'general'

                    words = content.strip().split()[:50]
                    preview = ' '.join(words) + ('...' if len(content.split()) > 50 else '')
                    print(f"Document preview: {preview}")

                if not raw_docs:
                    raise Exception(f"No content extracted from {file_type} file")

                # Update metadata with document type
                for doc in raw_docs:
                    doc.metadata['doc_type'] = doc_type

                # Filter minimal content
                filtered_docs = [
                    doc for doc in raw_docs
                    if len(doc.page_content.strip()) > 20
                    and doc.page_content.strip().count(' ') > 3
                ]

                if not filtered_docs:
                    filtered_docs = raw_docs

                # Calculate total content length for smart splitting
                total_length = sum(len(doc.page_content) for doc in filtered_docs)

                # Use universal splitter
                splitter = get_smart_splitter(total_length, doc_type)
                split_docs = splitter.split_documents(filtered_docs)

                # Filter chunks
                semantic_filtered = [
                    doc for doc in split_docs
                    if len(doc.page_content.strip()) > 10 and len(doc.page_content.split()) > 2
                ]

                return semantic_filtered

            except Exception as e:
                error_details = traceback.format_exc()
                raise Exception(f"Document processing failed for {file_path}: {str(e)}") from e

        split_docs = await run_in_threadpool(load_and_split)

        if not split_docs:
            raise Exception("No content extracted from document")

        # Apply semantic similarity-based sampling
        final_docs = await semantic_chunk_sampling(split_docs, max_chunks, query)

        return final_docs

    finally:
        # Clean up temporary files
        if temp_download and file_path and file_path.exists():
            try:
                os.remove(file_path)
            except:
                pass


# Specific loader functions for different document types
async def load_pdf_semantic(path_or_url: str, max_chunks: int = 100, query: str = None) -> List[Document]:
    """Load PDF with semantic similarity-based chunking"""
    temp_download = False

    if path_or_url.startswith(("http://", "https://")):
        print(f"Processing PDF from URL: {path_or_url}")

        async with httpx.AsyncClient(
            timeout=httpx.Timeout(60.0),
            limits=httpx.Limits(max_connections=1, max_keepalive_connections=0)
        ) as client:
            response = await client.get(path_or_url, follow_redirects=True)
            response.raise_for_status()

            file_path = Path(tempfile.mkstemp(suffix=".pdf")[1])
            with open(file_path, "wb") as f:
                f.write(response.content)
            temp_download = True
    else:
        file_path = Path(path_or_url)
        if not file_path.exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")

    # Extract and print document preview
    doc_preview = await run_in_threadpool(extract_pdf_preview, str(file_path))
    print(f"PDF preview: {doc_preview}")

    def load_and_split():
        try:
            loader = PyPDFLoader(str(file_path))
            raw_docs = loader.load()

            if not raw_docs:
                raise Exception("PyPDFLoader returned no documents")

            # Update metadata
            for doc in raw_docs:
                doc.metadata['doc_type'] = 'insurance'

            # Filter documents
            filtered_docs = [
                doc for doc in raw_docs
                if len(doc.page_content.strip()) > 20
                and doc.page_content.strip().count(' ') > 3
            ]

            if not filtered_docs:
                filtered_docs = raw_docs

            # Use section-aware splitter
            total_length = sum(len(doc.page_content) for doc in filtered_docs)
            splitter = get_smart_splitter(total_length, 'insurance')
            split_docs = splitter.split_documents(filtered_docs)

            # Filter chunks
            semantic_filtered = [
                doc for doc in split_docs
                if len(doc.page_content.strip()) > 10 and len(doc.page_content.split()) > 2
            ]

            return semantic_filtered

        except Exception as e:
            error_details = traceback.format_exc()
            raise Exception(f"PDF processing failed for {file_path}: {str(e)}") from e

    split_docs = await run_in_threadpool(load_and_split)

    if temp_download and file_path.exists():
        try:
            os.remove(file_path)
        except:
            pass

    if not split_docs:
        raise Exception("No content extracted from PDF")

    # Apply semantic similarity-based sampling
    final_docs = await semantic_chunk_sampling(split_docs, max_chunks, query)
    return final_docs


async def load_docx_semantic(path_or_url: str, max_chunks: int = 100, query: str = None) -> List[Document]:
    """Load DOCX with semantic similarity-based chunking"""
    temp_download = False

    if path_or_url.startswith(("http://", "https://")):
        print(f"Processing DOCX from URL: {path_or_url}")

        async with httpx.AsyncClient(
            timeout=httpx.Timeout(60.0),
            limits=httpx.Limits(max_connections=1, max_keepalive_connections=0)
        ) as client:
            response = await client.get(path_or_url, follow_redirects=True)
            response.raise_for_status()

            file_path = Path(tempfile.mkstemp(suffix=".docx")[1])
            with open(file_path, "wb") as f:
                f.write(response.content)
            temp_download = True
    else:
        file_path = Path(path_or_url)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

    # Extract and print document preview
    doc_preview = await run_in_threadpool(extract_docx_preview, str(file_path))
    print(f"DOCX preview: {doc_preview}")

    def load_and_split():
        try:
            raw_docs = load_docx_document(str(file_path))

            if not raw_docs:
                raise Exception("No content extracted from DOCX")

            # Update metadata
            for doc in raw_docs:
                doc.metadata['doc_type'] = 'general'

            # Calculate total length and use appropriate splitter
            total_length = sum(len(doc.page_content) for doc in raw_docs)
            splitter = get_smart_splitter(total_length, 'general')
            split_docs = splitter.split_documents(raw_docs)

            # Filter chunks
            semantic_filtered = [
                doc for doc in split_docs
                if len(doc.page_content.strip()) > 10 and len(doc.page_content.split()) > 2
            ]

            return semantic_filtered

        except Exception as e:
            error_details = traceback.format_exc()
            raise Exception(f"DOCX processing failed for {file_path}: {str(e)}") from e

    split_docs = await run_in_threadpool(load_and_split)

    if temp_download and file_path.exists():
        try:
            os.remove(file_path)
        except:
            pass

    if not split_docs:
        raise Exception("No content extracted from DOCX")

    # Apply semantic similarity-based sampling
    final_docs = await semantic_chunk_sampling(split_docs, max_chunks, query)
    return final_docs


async def load_email_semantic(path_or_url: str, max_chunks: int = 100, query: str = None) -> List[Document]:
    """Load Email with semantic similarity-based chunking"""
    temp_download = False

    if path_or_url.startswith(("http://", "https://")):
        print(f"Processing Email from URL: {path_or_url}")

        async with httpx.AsyncClient(
            timeout=httpx.Timeout(60.0),
            limits=httpx.Limits(max_connections=1, max_keepalive_connections=0)
        ) as client:
            response = await client.get(path_or_url, follow_redirects=True)
            response.raise_for_status()

            file_path = Path(tempfile.mkstemp(suffix=".eml")[1])
            with open(file_path, "wb") as f:
                f.write(response.content)
            temp_download = True
    else:
        file_path = Path(path_or_url)
        if not file_path.exists():
            raise FileNotFoundError(f"Email file not found: {file_path}")

    # Extract and print email preview
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        email_content = f.read()
    doc_preview = extract_email_preview(email_content)
    print(f"Email preview: {doc_preview}")

    def load_and_split():
        try:
            raw_docs = load_email_document(str(file_path))

            if not raw_docs:
                raise Exception("No content extracted from email")

            # Update metadata
            for doc in raw_docs:
                doc.metadata['doc_type'] = 'email'

            # Calculate total length and use appropriate splitter
            total_length = sum(len(doc.page_content) for doc in raw_docs)
            splitter = get_smart_splitter(total_length, 'email')
            split_docs = splitter.split_documents(raw_docs)

            # Filter chunks
            semantic_filtered = [
                doc for doc in split_docs
                if len(doc.page_content.strip()) > 10 and len(doc.page_content.split()) > 2
            ]

            return semantic_filtered

        except Exception as e:
            error_details = traceback.format_exc()
            raise Exception(f"Email processing failed for {file_path}: {str(e)}") from e

    split_docs = await run_in_threadpool(load_and_split)

    if temp_download and file_path.exists():
        try:
            os.remove(file_path)
        except:
            pass

    if not split_docs:
        raise Exception("No content extracted from email")

    # Apply semantic similarity-based sampling
    final_docs = await semantic_chunk_sampling(split_docs, max_chunks, query)
    return final_docs


# Main API functions
async def load_pdf_ultra_fast(pdf_path: str) -> List[Document]:
    """Load PDF with page metadata preserved"""
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(pdf_path)
        documents = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            
            if text.strip():
                # Create document with page metadata
                documents.append(Document(
                    page_content=text,
                    metadata={"page": page_num + 1}  # 1-indexed page numbers
                ))
        
        doc.close()
        return documents
    except Exception as e:
        print(f"Error loading PDF: {e}")
        return []


async def load_document_ultra_fast(path_or_url: str, query: str = None) -> List[Document]:
    """Ultra fast universal document loading - main API function"""
    return await load_document_semantic(path_or_url, max_chunks=100, query=query)


def cleanup_temp_files(*file_paths):
    """Clean up temporary files"""
    for path in file_paths:
        if path and os.path.exists(path):
            try:
                os.remove(path)
            except:
                pass


# Additional utility functions
def validate_environment():
    """Validate that required environment variables and libraries are available"""
    issues = []

    # Check API key
    if not VOYAGE_API_KEY:
        issues.append("VOYAGE_API_KEY environment variable not set")

    # Check libraries
    if not PyPDF2:
        issues.append("PyPDF2 library not available - PDF processing may be limited")

    if not docx2txt and not DocxDocument:
        issues.append("No DOCX processing library available (docx2txt or python-docx)")

    if issues:
        print("Environment validation issues:")
        for issue in issues:
            print(f"  - {issue}")
        return False

    return True


def get_supported_formats():
    """Get list of supported document formats"""
    formats = ['PDF']

    if docx2txt or DocxDocument:
        formats.append('DOCX')

    formats.extend(['Email (.eml)', 'Text files'])

    return formats


async def batch_load_documents(file_paths: List[str], max_chunks_per_doc: int = 50,
                              query: str = None) -> Dict[str, List[Document]]:
    """Load multiple documents in batch"""
    results = {}

    for file_path in file_paths:
        try:
            print(f"Processing: {file_path}")
            docs = await load_document_semantic(file_path, max_chunks_per_doc, query)
            results[file_path] = docs
            print(f"Successfully loaded {len(docs)} chunks from {file_path}")
        except Exception as e:
            print(f"Failed to load {file_path}: {str(e)}")
            results[file_path] = []

    return results


# Example usage and testing
if __name__ == "__main__":
    import asyncio

    async def test_example():
        """Example usage of the enhanced RAG utils"""

        # Validate environment
        if not validate_environment():
            print("Please install required libraries and set environment variables")
            return

        print("Supported formats:", get_supported_formats())

        # Example file paths (replace with actual files)
        test_files = [
            "example.pdf",
            "example.docx",
            "example.eml"
        ]

        # Test universal loader
        for file_path in test_files:
            if os.path.exists(file_path):
                try:
                    print(f"\n--- Testing {file_path} ---")
                    docs = await load_document_ultra_fast(file_path, query="insurance coverage")
                    print(f"Loaded {len(docs)} chunks")

                    # Show first chunk
                    if docs:
                        print(f"First chunk preview: {docs[0].page_content[:200]}...")
                        print(f"Metadata: {docs[0].metadata}")

                except Exception as e:
                    print(f"Error loading {file_path}: {e}")

    # Run the test
    asyncio.run(test_example())